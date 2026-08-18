from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime

FONT = "Calibri"


def set_font(run, size=11, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)


def add_hyperlink(paragraph, text, url):
    if not url:
        return

    rid = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)

    run.append(rpr)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_section_heading(document, text, font_size):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0

    r = p.add_run(text.upper())
    set_font(r, font_size, bold=True)

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "5")
    bottom.set(qn("w:space"), "1")
    pBdr.append(bottom)
    pPr.append(pBdr)

    return p


def add_bullet(document, text):
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05

    r = p.add_run(str(text))
    set_font(r, 11)


def descriptions(value):
    if isinstance(value, str):
        return [value]
    if isinstance(value, list):
        return value
    return []


def generate_word_document(resume):
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(9)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)

    personal = resume.get("personal_information", {})

    full_name = personal.get("full_name", "")
    email = personal.get("email", "")
    phone = personal.get("phone", "")
    location = personal.get("location", "")
    linkedin = personal.get("linkedin", "")
    github = personal.get("github", "")
    portfolio = personal.get("portfolio", "")

    work_history = resume.get("work_history", [])

    current_role = ""
    for work in work_history:
        if work.get("currently_working"):
            current_role = work.get("role", "")
            break

    if not current_role and work_history:
        current_role = work_history[0].get("role", "")

    # ==========================================================
    # HEADER
    # ==========================================================

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)

    if full_name:
        r = p.add_run(full_name.upper())
        set_font(r, 16, bold=True)

    if current_role:
        r = p.add_run(" | ")
        set_font(r, 16, bold=True)

        r = p.add_run(current_role.upper())
        set_font(r, 16, bold=True)

    contact = []

    if email:
        contact.append(f"Email: {email}")
    if phone:
        contact.append(f"Phone: {phone}")
    if location:
        contact.append(f"Address: {location}")

    if contact:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)

        r = p.add_run(" | ".join(contact))
        set_font(r, 11)

    links = [
        ("LinkedIn", linkedin),
        ("GitHub", github),
        ("Portfolio", portfolio)
    ]

    links = [(name, url) for name, url in links if url]

    if links:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)

        for i, (label, url) in enumerate(links):
            if i:
                r = p.add_run(" | ")
                set_font(r, 11)
            add_hyperlink(p, label, url)

    # ==========================================================
    # SUMMARY
    # ==========================================================

    summary = resume.get("professional_summary", "")
    if summary:
        add_section_heading(document, "Professional Summary", font_size=12)
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(summary)
        set_font(r, 11)

    # ==========================================================
    # SKILLS
    # ==========================================================

    skills = resume.get("skills", {})
    if skills:
        add_section_heading(document, "Skills", font_size=12)

        for category, skill_list in skills.items():
            if not skill_list:
                continue

            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0

            r = p.add_run(f"{category}: ")
            set_font(r, 11, bold=True)

            r = p.add_run(", ".join(skill_list))
            set_font(r, 11)

    # ==========================================================
    # WORK HISTORY
    # ==========================================================

    if work_history:
        add_section_heading(document, "Work History", font_size=12)

        for work in work_history:
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(0)

            r = p.add_run(work.get("role", ""))
            set_font(r, 11, bold=True)

            company = work.get("company", "")
            if company:
                r = p.add_run(f" | {company}")
                set_font(r, 11, bold=True)

            start_date = work.get("start_date", "")

            if start_date:
                start = datetime.strptime(
                    start_date,
                    "%Y-%m-%d"
                ).strftime("%b %Y")
            else:
                start = ""

            if work.get("currently_working"):
                end = "Present"
            else:
                end_date = work.get("end_date", "")

                if end_date:
                    end = datetime.strptime(
                        end_date,
                        "%Y-%m-%d"
                    ).strftime("%b %Y")
                else:
                    end = ""

            if start or end:
                p = document.add_paragraph()
                p.paragraph_format.space_after = Pt(2)

                r = p.add_run(f"{start} – {end}")
                set_font(r, 11, italic=True)

            for item in descriptions(work.get("description", [])):
                if item:
                    add_bullet(document, item)

    # ==========================================================
    # PROJECTS
    # ==========================================================

    def add_projects(title, projects):
        if not projects:
            return

        add_section_heading(document, title, font_size=12)

        for project in projects:
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(0)

            r = p.add_run(project.get("title", ""))
            set_font(r, 11, bold=True)

            project_github = project.get("github", "")
            if project_github:
                r = p.add_run(" | ")
                set_font(r, 11)
                add_hyperlink(p, "GitHub", project_github)

            technologies = project.get("technologies", [])
            if technologies:
                p = document.add_paragraph()
                p.paragraph_format.space_after = Pt(2)

                r = p.add_run("Technologies: ")
                set_font(r, 11, bold=True)

                r = p.add_run(", ".join(technologies))
                set_font(r, 11)

            for item in descriptions(project.get("description", [])):
                if item:
                    add_bullet(document, item)

    add_projects("Professional Projects", resume.get("professional_projects", []))
    add_projects("Academic Projects", resume.get("academic_projects", []))
    add_projects("Personal Projects", resume.get("personal_projects", []))

    # ==========================================================
    # EDUCATION
    # ==========================================================

    education = resume.get("education", [])

    if education:
        add_section_heading(document, "Education", font_size=12)

        for edu in education:
            level = edu.get("level", "")
            specialization = edu.get("specialization", "")
            institution = edu.get("institution", "")
            board = edu.get("board_university", "")
            start = edu.get("start_year", "")
            end = edu.get("end_year", "")
            score = edu.get("score", "")
            score_type = edu.get("score_type", "")

            level_lower = level.lower()

            if level_lower == "bachelor's":
                degree = "Bachelor of Technology"
            elif level_lower == "intermediate":
                degree = "Intermediate"
            elif level_lower == "ssc":
                degree = "Class 10th"
            else:
                degree = level

            title = f"{degree} ({specialization})" if specialization else degree

            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(0)

            r = p.add_run(title)
            set_font(r, 11, bold=True)

            if board and level_lower != "bachelor's":
                r = p.add_run(f", {board}")
                set_font(r, 11)

            if institution:
                p = document.add_paragraph()
                p.paragraph_format.space_after = Pt(0)

                text = f"Institute: {institution}"
                if board and level_lower == "bachelor's":
                    text += f", affiliated with {board}"

                r = p.add_run(text)
                set_font(r, 11)

            details = []

            if score:
                if score_type:
                    score_text = f"{score} {score_type}"
                else:
                    score_text = str(score)

                details.append(f"Grade Points: {score_text}")

            if start and end:
                details.append(f"{start}-{end}")

            if details:
                p = document.add_paragraph()
                p.paragraph_format.space_after = Pt(2)

                r = p.add_run(", ".join(details))
                set_font(r, 11)

    # ==========================================================
    # CERTIFICATIONS
    # ==========================================================

    certifications = resume.get("certifications", [])

    if certifications:
        add_section_heading(document, "Certifications", font_size=12)

        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05

        r = p.add_run("Certifications: ")
        set_font(r, 11, bold=True)

        values = []

        for cert in certifications:
            name = cert.get("name", "")
            issuer = cert.get("issuer", "")

            if not name:
                continue

            values.append(
                f"{name} by {issuer}" if issuer else name
            )

        r = p.add_run(", ".join(values))
        set_font(r, 11)

    # ==========================================================
    # Final spacing
    # ==========================================================

    for p in document.paragraphs:
        if p.paragraph_format.line_spacing is None:
            p.paragraph_format.line_spacing = 1.0

    output = BytesIO()
    document.save(output)
    output.seek(0)

    return output.getvalue()